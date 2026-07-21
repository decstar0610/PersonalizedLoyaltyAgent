"""Generate a 2-page project summary report (LoyaltyForge / PS100) as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NV_GREEN = RGBColor(0x76, 0xB9, 0x00)
INK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9)
style.font.color.rgb = INK
pf = style.paragraph_format
pf.space_after = Pt(2.5)
pf.line_spacing = 1.0

for section in doc.sections:
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NV_GREEN
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1.5)
    return p


def body(text):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def bullet(lead, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.left_indent = Inches(0.2)
    if lead:
        r = p.add_run(lead + " ")
        r.bold = True
    p.add_run(text)
    return p


# ---------- Title ----------
title = doc.add_paragraph()
tr = title.add_run("LoyaltyForge — Personalized Reward Journeys")
tr.bold = True
tr.font.size = Pt(16)
tr.font.color.rgb = INK
title.paragraph_format.space_after = Pt(0)

sub = doc.add_paragraph()
sr = sub.add_run("Project Summary Report  ·  NVIDIA Final-Project Submission (PS100)  ·  Agentic AI · Retail & E-commerce")
sr.font.size = Pt(9)
sr.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(0.5)

meta = doc.add_paragraph()
mr = meta.add_run("Author: Priyanka M    |    Domain: Retail & E-commerce    |    Stack: End-to-end NVIDIA NIM    |    Version 1.1")
mr.font.size = Pt(8)
mr.italic = True
mr.font.color.rgb = GREY
meta.paragraph_format.space_after = Pt(3)

divp = doc.add_paragraph()
divr = divp.add_run("_" * 130)
divr.font.color.rgb = NV_GREEN
divr.font.size = Pt(4)
divp.paragraph_format.space_after = Pt(1)

# ---------- 1 ----------
heading("1 · Problem Statement & Goal")
body("Retailers run generic, one-size-fits-all loyalty programs: every customer in a tier gets the "
     "same offers, producing irrelevant promotions, low redemption, and churn — and these blanket "
     "campaigns frequently ignore each customer's privacy and consent boundaries. The goal of "
     "LoyaltyForge is to replace that static rulebook with an autonomous agent that reasons about "
     "each customer individually and produces a tailored loyalty reward journey, while treating "
     "consent as a hard, non-negotiable constraint. The system runs end-to-end on the NVIDIA NIM "
     "stack — NVIDIA was both the platform requirement and the reasoning engine.")

# ---------- 2 ----------
heading("2 · Solution Overview — Three Agentic Behaviors")
body("Given a customer_id the agent reads purchase history, tier and points; checks consent BEFORE "
     "any personalization; retrieves applicable loyalty rules via RAG; and reasons over all of it to "
     "emit a structured, schema-validated journey. Three behaviors make it genuinely agentic rather "
     "than a single LLM call:")
bullet("Dynamic tool routing —", "LangGraph conditional edges classify each customer (new, lapsing, "
       "near a tier threshold, established) and route to a segment-specific rule path. Consent is the first gate.")
bullet("Multi-flag consent reasoning —", "personalization off → generic offer, stop. When on, it honors "
       "secondary flags: email_marketing off → no email channel; data_sharing off → first-party data only.")
bullet("Self-critique loop —", "a validation node checks reward grounding, consent compliance, and schema; "
       "on failure the agent revises and re-validates (≤2 retries) before finalizing.")

# ---------- 3 ----------
heading("3 · Architecture & Agent Flow")
body("USER (React Studio / API) → CONSENT GATE (off → generic offer, STOP) → CLASSIFY + ROUTE "
     "(new / lapsing / near-threshold / established) → segment-specific rules via FAISS RAG (NIM "
     "embeddings) → DRAFT (NVIDIA NIM LLM, honoring consent) → VALIDATE ⇄ REVISE (grounding · consent · "
     "schema, ≤2 retries) → FINALIZE (Pydantic-enforced) → Journey JSON back to the Studio / Streamlit.")

# ---------- 4 ----------
heading("4 · Dataset & Data Model")
body("All data sources share one schema (customer_id, name, persona, purchase_history [item, category, "
     "amount, date], consent_flags, loyalty_tier, points), so the agent runs on any of them unchanged via "
     "the CUSTOMERS_FILE env var. Loyalty rules live separately in data/loyalty_rules.md — the RAG source "
     "embedded into FAISS with NIM embeddings and retrieved per segment to ground rewards.")
bullet("Curated demo set (default) —", "data/customers.json: 5 hand-built customers chosen to exercise "
       "every path — C001 Asha (silver, data_sharing off), C002 Rahul (gold, all consent on), C003 Meera "
       "(bronze, new, email off), C004 Vikram (personalization OFF → hard consent-gate case, no LLM call), "
       "C005 Sofia (platinum, broad/established). Small by design: enough to demo all behaviors deterministically.")
bullet("Real-world scale-up —", "Kaggle “Customer Personality Analysis” (imakash3011, marketing_campaign.csv, "
       "~2,240 rows, one row per customer) — picked over UCI Online Retail II for cleaner fit (already per-customer). "
       "Ingestion script maps Mnt* category spend → purchase_history; synthesizes points/tier from total spend "
       "(Bronze 0 / Silver 250 / Gold 500 / Platinum 1000) and persona from demographics. Two fields are "
       "synthesized (no dataset has them): anonymized name and consent_flags.personalization (deterministic, "
       "so some members stay consent-off for the demo).")
bullet("Synthetic generator —", "scripts/generate_customers.py --count N produces arbitrary-size sets; a "
       "200-row data/customers_large.json was used to prove scaling. The curated 5 remain the default.")

# ---------- 5 ----------
heading("5 · Key Design Choices & Rationale")
bullet("NVIDIA NIM (LLM + embeddings) —", "meta/llama-3.1-70b-instruct + nvidia/nv-embedqa-e5-v5, keeping "
       "the stack fully NVIDIA (the core requirement) rather than mixing providers.")
bullet("LangGraph / LangChain —", "real dynamic routing via conditional edges and a clean home for the validate→revise loop.")
bullet("FAISS + Pydantic —", "lightweight in-process RAG grounds rewards in the rules doc; Pydantic schema-validates "
       "the whole response so the UI/API always get valid JSON.")
bullet("Consent-first hard gate —", "a consent-off customer is handled deterministically and never even calls the "
       "LLM — the headline differentiator and a responsible-AI choice.")
bullet("FastAPI + React “Customer Journey Studio” —", "POST /generate-journey, GET /customers, GET /health; the "
       "premium React/Vite/Tailwind Studio is the primary demo, with Streamlit as a simple backup.")
bullet("Honest, derived metrics —", "Studio “Loyalty DNA” and “Business Impact” numbers are computed "
       "deterministically from real data and labelled “estimate”; only the journey itself comes from the agent.")

# ---------- 6 tech table ----------
heading("6 · Technology Stack")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("Layer").bold = True
hdr[1].paragraphs[0].add_run("Technology").bold = True
rows = [
    ("LLM / Embeddings", "NVIDIA NIM — meta/llama-3.1-70b-instruct / nvidia/nv-embedqa-e5-v5"),
    ("Agent framework / Schema", "LangGraph + LangChain / Pydantic"),
    ("Vector store / Backend", "FAISS (faiss-cpu) / FastAPI + Uvicorn"),
    ("Studio UI / Backup UI", "React + Vite + Tailwind / Streamlit"),
    ("Data / Language", "JSON + Markdown (swappable via CUSTOMERS_FILE) / Python 3.10+ & TypeScript"),
]
for layer, tech in rows:
    c = tbl.add_row().cells
    rp = c[0].paragraphs[0].add_run(layer); rp.font.size = Pt(8.3); rp.bold = True
    rt = c[1].paragraphs[0].add_run(tech); rt.font.size = Pt(8.3)
for row in tbl.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)

# ---------- 7 ----------
heading("7 · Testing & Reliability")
body("The agent's safety-critical logic is covered by 17 offline pytest tests (all green). The NVIDIA "
     "LLM and RAG retrieval are mocked, so the suite is fast, deterministic, and needs no API key or "
     "network — its job is to catch regressions in the guardrails. It covers consent gating (consent-off "
     "→ generic offer only), segment routing (correct new/lapsing/near-threshold/established classification), "
     "multi-flag consent constraints (email and data-sharing restrictions), and the self-critique loop "
     "(validator flags ungrounded rewards, consent violations, and schema errors; the agent revises before "
     "finalizing). The tests verify deterministic guardrails and schema — not the LLM's wording.")

# ---------- 8 ----------
heading("8 · Status, Scope & Next Steps")
bullet("Done —", "consent-gated NVIDIA agent with all three behaviors, React Studio wired to FastAPI, 17 "
       "passing offline tests, fully updated PRD + README; pushed to a private GitHub repo (main, in sync).")
bullet("Out of scope (this version) —", "production database, payment/rewards fulfillment, user auth, "
       "A/B-testing pipelines — mock JSON data is sufficient for the MVP.")
bullet("Open items —", "record the demo video / screenshots (only remaining PRD deliverable); optional Kaggle "
       "scale-up (pipeline built, awaiting the CSV); optionally surface segment/validation_passed/consent_notes in the UI.")
bullet("Future —", "Postgres + live transaction feeds, multi-agent split, reward fulfillment, redemption-rate "
       "feedback loop, multilingual messages.")

foot = doc.add_paragraph()
fr = foot.add_run("Success criteria (target 100% / Yes): a valid journey for every customer; a consent-off "
     "customer never receives personalized offers (hard requirement); valid JSON matching the Pydantic schema; "
     "the full test suite passes; the end-to-end demo runs without crashing; the README lets a stranger set it up.")
fr.font.size = Pt(8)
fr.italic = True
fr.font.color.rgb = GREY
foot.paragraph_format.space_before = Pt(3)

out = r"D:\projects\projects\PersonalizedLoyaltyAgent-nvidia\docs\LoyaltyForge_Project_Summary.docx"
doc.save(out)
print("Saved:", out)
